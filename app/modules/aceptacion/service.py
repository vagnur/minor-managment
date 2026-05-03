import re
from pathlib import Path

from docx import Document

from app.core.excel_utils import read_excel_file
from app.core.file_utils import ensure_folder, sanitize_filename, safe_str
from app.core.validation_utils import (
    validate_file_exists,
    validate_required_columns,
    validate_non_empty_dataframe,
)
from app.core.docx_utils import (
    find_table_by_headers,
    clone_last_row,
    clear_row_text,
    set_cell_no_wrap,
    set_table_fixed_layout,
)


REQUIRED_INTERNAL_COLUMNS = [
    "RUT",
    "NombreEstudiante",
    "Carrera",
    "Facultad",
]


def normalize_dataframe(df, column_mapping: dict):
    return df.rename(columns=column_mapping)


def normalize_rut(rut: str) -> tuple[str, str]:
    rut = safe_str(rut).upper().replace(".", "").strip()

    if "-" not in rut:
        raise ValueError(f"RUT sin guión: {rut}")

    run, dv = rut.split("-", 1)
    run = run.strip()
    dv = dv.strip()

    if not run.isdigit():
        raise ValueError(f"RUN inválido: {rut}")

    if not re.fullmatch(r"[0-9K]", dv):
        raise ValueError(f"DV inválido: {rut}")

    return run, dv


def normalize_facultad(facultad: str) -> str:
    facultad = safe_str(facultad).strip()

    replacements = {
        "Facultad de Ingeniería": "Ingeniería",
        "Ingeniería": "Ingeniería",
    }

    return replacements.get(facultad, facultad)


def format_semestre_texto(semestre: str, anio: str) -> str:
    semestre = safe_str(semestre).strip()

    mapping = {
        "1": "Primer Semestre de",
        "2": "Segundo Semestre de",
    }

    prefix = mapping.get(semestre, f"Semestre {semestre} de")
    return f"{prefix} {safe_str(anio)}"


def split_full_name(full_name: str) -> tuple[str, str, str | None]:
    full_name = " ".join(safe_str(full_name).split())
    parts = full_name.split()

    if len(parts) not in (3, 4):
        raise ValueError(
            f"Nombre con formato inválido: '{full_name}'. "
            "Se esperaba 1 o 2 nombres y 2 apellidos."
        )

    apellidos = " ".join(parts[-2:])
    nombres = " ".join(parts[:-2])
    return apellidos, nombres, None


def build_output_filename(config: dict, semestre: str, anio: str) -> str:
    pattern = config["output_filename_pattern"]
    filename = pattern.format(semestre=semestre, anio=anio)
    return sanitize_filename(filename.replace(".docx", "")) + ".docx"


def load_acceptance_dataframe(excel_path: str, config: dict):
    df = read_excel_file(excel_path, sheet_name=0)
    df = normalize_dataframe(df, config["column_mapping"])
    validate_non_empty_dataframe(df)
    validate_required_columns(df, REQUIRED_INTERNAL_COLUMNS)
    return df


def validate_template_structure(doc, config: dict):
    expected_headers = config["table_headers"]
    table = find_table_by_headers(doc, expected_headers)

    if table is None:
        raise ValueError(
            "No se encontró la tabla esperada en la plantilla con los encabezados configurados."
        )

    required_markers = [
        "SEMESTRE_INGRESO",
        "AÑO_SEMESTRE",
        "INICIALES_DIRECTOR_DEPA",
        "INICIALES_COORDINADOR_MINOR",
    ]

    full_text = "\n".join(p.text for p in doc.paragraphs)
    full_text += "\n"
    for table_ in doc.tables:
        for row in table_.rows:
            for cell in row.cells:
                full_text += cell.text + "\n"

    missing_markers = [m for m in required_markers if m not in full_text]
    if missing_markers:
        raise ValueError(
            "Faltan marcadores en la plantilla: " + ", ".join(missing_markers)
        )

    return table


def replace_text_in_run_preserving_format(doc, replacements: dict):
    """
    Reemplaza placeholders recorriendo runs para no perder formato
    (negritas, cursivas, etc.).
    """
    def replace_in_paragraph(paragraph):
        for run in paragraph.runs:
            original = run.text
            updated = original
            for old, new in replacements.items():
                if old in updated:
                    updated = updated.replace(old, new)
            if updated != original:
                run.text = updated

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def fill_acceptance_table(table, df, logger=None):
    warnings = []

    set_table_fixed_layout(table)

    for idx, (_, row) in enumerate(df.iterrows(), start=1):
        rut = safe_str(row["RUT"])
        nombre = safe_str(row["NombreEstudiante"])
        carrera = safe_str(row["Carrera"])
        facultad = normalize_facultad(row["Facultad"])

        run, dv = normalize_rut(rut)
        apellidos, nombres, warning = split_full_name(nombre)

        if warning:
            warnings.append(f"Fila {idx}: {warning}")
            if logger:
                logger(f"Advertencia fila {idx}: {warning}")

        new_row = clone_last_row(table)
        clear_row_text(new_row)

        new_row.cells[0].text = str(idx)
        new_row.cells[1].text = run
        new_row.cells[2].text = dv
        new_row.cells[3].text = apellidos
        new_row.cells[4].text = nombres
        new_row.cells[5].text = carrera
        new_row.cells[6].text = facultad

        set_cell_no_wrap(new_row.cells[0])
        set_cell_no_wrap(new_row.cells[1])
        set_cell_no_wrap(new_row.cells[2])

    if len(table.rows) > 1:
        last_row = table.rows[-1]
        row_text = "".join(cell.text.strip() for cell in last_row.cells)
        if row_text == "":
            tbl = table._tbl
            tbl.remove(last_row._tr)

    return warnings


def process_aceptacion(
    excel_path: str,
    output_folder: str,
    semestre: str,
    anio: str,
    iniciales_director: str,
    iniciales_coordinador: str,
    config: dict,
    logger=None,
):
    def log(msg: str):
        if logger:
            logger(msg)

    validate_file_exists(excel_path, "archivo Excel")
    validate_file_exists(config["template_path"], "plantilla Word")

    df = load_acceptance_dataframe(excel_path, config)
    log(f"Registros cargados: {len(df)}")

    doc = Document(config["template_path"])
    table = validate_template_structure(doc, config)

    replacements = {
        "SEMESTRE_INGRESO": format_semestre_texto(semestre, anio),
        "AÑO_SEMESTRE": safe_str(anio),
        "INICIALES_DIRECTOR_DEPA": safe_str(iniciales_director),
        "INICIALES_COORDINADOR_MINOR": safe_str(iniciales_coordinador),
    }

    replace_text_in_run_preserving_format(doc, replacements)

    warnings = fill_acceptance_table(table, df, logger=log)

    ensure_folder(output_folder)
    output_name = build_output_filename(config, semestre, anio)
    output_path = Path(output_folder) / output_name

    doc.save(output_path)

    return {
        "total": len(df),
        "warnings": warnings,
        "output_path": str(output_path),
    }