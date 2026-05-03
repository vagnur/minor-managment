from pathlib import Path

import pandas as pd
from docx import Document

from app.core.file_utils import ensure_folder, sanitize_filename, safe_str
from app.core.validation_utils import validate_file_exists, validate_required_columns

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from app.core.docx_utils import clone_last_row, clear_row_text, set_table_fixed_layout, set_cell_no_wrap

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def normalize_text(text: str) -> str:
    return (
        safe_str(text)
        .lower()
        .replace(" ", "")
        .replace("_", "")
        .replace("-", "")
        .replace(".", "")
    )

def write_cell(cell, value):
    cell.text = ""

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = paragraph.add_run(safe_str(value))
    run.bold = False  # 👈 importante

    set_cell_border(cell)

def detect_subject_from_filename(filename: str, config: dict) -> str | None:
    normalized_filename = normalize_text(filename)

    matches = []

    for subject_name, subject_config in config["subjects"].items():
        aliases = subject_config.get("aliases", [])

        for alias in aliases:
            normalized_alias = normalize_text(alias)

            if normalized_alias in normalized_filename:
                matches.append((subject_name, len(normalized_alias)))

    if not matches:
        return None

    matches.sort(key=lambda item: item[1], reverse=True)
    return matches[0][0]


def clean_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df = df.loc[:, ~df.columns.astype(str).str.startswith("Unnamed")]
    df.columns = [str(col).strip() for col in df.columns]
    df = df.dropna(how="all")

    return df


def read_excel_file(file_path: Path, subject_name: str, config: dict) -> pd.DataFrame:
    df = pd.read_excel(file_path, sheet_name=0)
    df = clean_dataframe(df)

    subject_config = config["subjects"][subject_name]
    required_columns = subject_config.get("required_columns", [])

    validate_required_columns(df, required_columns)

    return df


def read_grades_folder(folder_path: str, config: dict, logger=None) -> list[dict]:
    def log(message: str):
        if logger:
            logger(message)

    folder = Path(folder_path)

    if not folder.exists():
        raise FileNotFoundError(f"No existe la carpeta seleccionada: {folder}")

    if not folder.is_dir():
        raise NotADirectoryError(f"La ruta seleccionada no corresponde a una carpeta: {folder}")

    excel_files = sorted([
        file_path
        for file_path in folder.iterdir()
        if file_path.suffix.lower() in [".xlsx", ".xls"]
        and not file_path.name.startswith("~$")
    ])

    if not excel_files:
        raise ValueError("La carpeta seleccionada no contiene archivos Excel.")

    grade_files = []

    for file_path in excel_files:
        subject_name = detect_subject_from_filename(file_path.name, config)

        if subject_name is None:
            log(f"[OMITIDO] No se pudo detectar asignatura: {file_path.name}")
            continue

        try:
            df = read_excel_file(file_path, subject_name, config)

            grade_files.append({
                "file_path": file_path,
                "file_name": file_path.name,
                "subject": subject_name,
                "subject_config": config["subjects"][subject_name],
                "dataframe": df,
                "student_count": len(df),
            })

            log(f"[OK] {file_path.name} → {subject_name} → {len(df)} estudiantes")

        except Exception as e:
            log(f"[ERROR] {file_path.name}: {e}")
            raise

    if not grade_files:
        raise ValueError("No se pudo procesar ningún archivo de notas válido.")

    return grade_files


def preview_grades_folder(folder_path: str, config: dict, logger=None) -> dict:
    grade_files = read_grades_folder(
        folder_path=folder_path,
        config=config,
        logger=logger,
    )

    subjects_found = sorted(set(item["subject"] for item in grade_files))
    total_students = sum(item["student_count"] for item in grade_files)

    return {
        "files_read": len(grade_files),
        "subjects_found": subjects_found,
        "total_students": total_students,
        "grade_files": grade_files,
    }

def generate_faculty_excels(
    folder_path: str,
    output_folder: str,
    config: dict,
    logger=None,
) -> dict:

    def log(msg: str):
        if logger:
            logger(msg)

    grade_files = read_grades_folder(folder_path, config, logger)

    faculty_data = {}

    # 🔹 Agrupar datos
    for item in grade_files:
        df = item["dataframe"]
        subject = item["subject"]

        if "Facultad" not in df.columns:
            raise ValueError(f"El archivo {item['file_name']} no tiene columna 'Facultad'")

        for facultad in df["Facultad"].dropna().unique():
            df_facultad = df[df["Facultad"] == facultad]

            if facultad not in faculty_data:
                faculty_data[facultad] = {}

            if subject not in faculty_data[facultad]:
                faculty_data[facultad][subject] = df_facultad
            else:
                faculty_data[facultad][subject] = pd.concat(
                    [faculty_data[facultad][subject], df_facultad],
                    ignore_index=True
                )

    # 🔹 Crear archivos Excel
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)

    total_files = 0

    for facultad, subjects_dict in faculty_data.items():
        file_name = f"{facultad}.xlsx".replace("/", "-")
        file_path = output_path / file_name

        with pd.ExcelWriter(file_path, engine="openpyxl") as writer:
            for subject, df_subject in subjects_dict.items():
                df_subject.to_excel(writer, sheet_name=subject, index=False)

        total_files += 1
        log(f"[OK] Excel generado: {file_path}")

    return {
        "total_faculties": total_files,
        "faculties": list(faculty_data.keys()),
    }

def replace_placeholders_in_paragraphs(doc, replacements: dict):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, safe_str(value))


def replace_placeholders_in_tables(doc, replacements: dict):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, safe_str(value))


def fill_memo_table_regular(doc, df: pd.DataFrame):
    table = doc.tables[0]
    set_table_fixed_layout(table)

    for idx, (_, row) in enumerate(df.iterrows(), start=1):
        new_row = clone_last_row(table)
        clear_row_text(new_row)

        cells = new_row.cells

        write_cell(cells[0], idx)
        write_cell(cells[1], row["Nombre"])
        write_cell(cells[2], row["RUT Estudiante"])
        write_cell(cells[3], row["Carrera"])
        write_cell(cells[4], row["Sección Laboratorio"])
        write_cell(cells[5], row["Nota Cátedra"])
        write_cell(cells[6], row["Nota Laboratorio"])
        write_cell(cells[7], row["Promedio"])


def fill_memo_table_taaa(doc, df: pd.DataFrame):
    table = doc.tables[0]
    set_table_fixed_layout(table)

    for idx, (_, row) in enumerate(df.iterrows(), start=1):
        new_row = clone_last_row(table)
        clear_row_text(new_row)

        cells = new_row.cells

        write_cell(cells[0], idx)
        write_cell(cells[1], row["Nombre"])
        write_cell(cells[2], row["RUT Estudiante"])
        write_cell(cells[3], row["Carrera"])
        write_cell(cells[4], row["Sección Laboratorio"])
        write_cell(cells[5], row["Nota Laboratorio"])
        write_cell(cells[6], row["Promedio"])


def build_memo_output_name(subject_name: str, section_name: str, professor_name: str) -> str:
    return sanitize_filename(
        f"memorandum_{subject_name}_{section_name}_{professor_name}.docx"
    )

def extract_year_from_fecha(fecha: str) -> str:
    parts = safe_str(fecha).strip().split()

    for part in reversed(parts):
        if part.isdigit() and len(part) == 4:
            return part

    raise ValueError(
        "No se pudo obtener el año desde la fecha. "
        "Usa formato tipo: '16 de marzo de 2026'."
    )


def generate_single_memo(
    df: pd.DataFrame,
    subject_name: str,
    subject_config: dict,
    fecha: str,
    vice: str,
    semestre: str,
    output_folder: str,
) -> Path:
    is_taaa = not subject_config.get("has_catedra", True)
    anio = extract_year_from_fecha(fecha)

    template_path = subject_config["template_path"]
    validate_file_exists(template_path, f"plantilla de notas para {subject_name}")

    if is_taaa:
        professor_name = safe_str(df.iloc[0]["Profesor Laboratorio"])
        professor_rut = safe_str(df.iloc[0]["RUT Profesor Laboratorio"])
        section_name = safe_str(df.iloc[0]["Sección Laboratorio"])
    else:
        professor_name = safe_str(df.iloc[0]["Profesor Cátedra"])
        professor_rut = safe_str(df.iloc[0]["RUT Profesor Cátedra"])
        section_name = safe_str(df.iloc[0]["Sección Cátedra"])

    doc = Document(template_path)

    replacements = {
        "<anio>": anio,
        "<fecha>": fecha,
        "<vice>": vice,
        "<asignatura>": subject_config["display_name"],
        "<codigo>": subject_config["code"],
        "<semestre>": semestre,
        "<profesor>": professor_name,
        "<rut_profe>": professor_rut,
    }

    replace_placeholders_in_paragraphs(doc, replacements)
    replace_placeholders_in_tables(doc, replacements)

    if is_taaa:
        fill_memo_table_taaa(doc, df)
    else:
        fill_memo_table_regular(doc, df)

    output_path = ensure_folder(Path(output_folder) / "memorandums")
    file_name = build_memo_output_name(subject_name, section_name, professor_name)
    save_path = output_path / file_name

    doc.save(save_path)

    return save_path


def generate_memos(
    folder_path: str,
    output_folder: str,
    fecha: str,
    vice: str,
    semestre: str,
    config: dict,
    logger=None,
) -> dict:
    def log(msg: str):
        if logger:
            logger(msg)

    extract_year_from_fecha(fecha)

    grade_files = read_grades_folder(folder_path, config, logger)

    result = {
        "total_ok": 0,
        "total_errors": 0,
        "generated_files": [],
        "error_details": [],
    }

    for item in grade_files:
        subject_name = item["subject"]
        subject_config = item["subject_config"]
        df = item["dataframe"]

        is_taaa = not subject_config.get("has_catedra", True)

        if is_taaa:
            group_column = "Sección Laboratorio"
        else:
            group_column = "Sección Cátedra"

        log(f"[{subject_name}] Generando memorándums por {group_column}...")

        for section_value, df_group in df.groupby(group_column, dropna=False):
            try:
                save_path = generate_single_memo(
                    df=df_group,
                    subject_name=subject_name,
                    subject_config=subject_config,
                    fecha=fecha,
                    vice=vice,
                    semestre=semestre,
                    output_folder=output_folder,
                )

                result["total_ok"] += 1
                result["generated_files"].append(str(save_path))
                log(f"[OK] Memo generado: {save_path}")

            except Exception as e:
                result["total_errors"] += 1
                detail = f"[{subject_name}] Error generando memo sección {section_value}: {e}"
                result["error_details"].append(detail)
                log(f"[ERROR] {detail}")

    return result

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for edge in ('top', 'left', 'bottom', 'right'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '8')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        tcPr.append(element)