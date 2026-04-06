from pathlib import Path
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.file_utils import ensure_folder, sanitize_filename, safe_str
from app.core.excel_utils import read_excel_file
from app.core.validation_utils import (
    validate_file_exists,
    validate_required_columns,
    validate_non_empty_dataframe,
)
from app.core.docx_utils import find_table_by_text, find_row_index_by_cell_text


COLUMN_RENAME_MAP = {
    "Marca temporal": "Fecha",
    "Dirección de correo electrónico": "Correo",
    #"Minor al que postula": "Minor",
    "Primer Nombre": "PrimerNombre",
    "Segundo Nombre": "SegundoNombre",
    "Apellido paterno": "ApellidoPaterno",
    "Apellido materno": "ApellidoMaterno",
    "RUT": "RUT",
    #"Número de celular o de contacto": "NumeroCelular",
    "Correo institucional": "CorreoInstitucional",
    #"Correo personal (diferente al institucional)": "CorreoPersonal",
    "Carrera a la que pertenece": "Carrera",
    "Nombre y apellido de su Jefe Carrera": "JefeCarrera",
    "Correo electrónico de su Jefe de Carrera": "CorreoJefeCarrera",
    "Duración de la carrera": "DuracionCarrera",
    "Avance curricular": "AvanceCurricular",
    "Facultad a la que pertenece": "Facultad",
    #"Indique asignatura a la cual postula": "AsignaturaPostulada",
    "Seleccione el o los horarios de cátedra a los cuales puede asistir": "HorariosCatedra",
    "Seleccione el o los horarios de laboratorio a los cuales puede asistir": "HorariosLaboratorio",
    "Indique todos los horarios disponibles en los cuales podría participar en la asignatura, esto es para hacer un catastro entre aquellas/os estudiantes que no pueden participar en el Minor debido a topes de horarios. Nos permite ver posibilidades de solicitar una nueva coordinación en caso de haber demanda suficiente": "HorariosDisponibles",
    "¿Por qué quiere hacer este Minor?": "MotivoMinor",
    "¿Qué espera aprender en este Minor?": "ExpectativasMinor",
    "Comentarios adicionales": "Comentarios"
}


REQUIRED_COLUMNS = [
    "PrimerNombre",
    "SegundoNombre",
    "ApellidoPaterno",
    "ApellidoMaterno",
    "RUT",
    "CorreoInstitucional",
    "Carrera",
    "JefeCarrera",
    "DuracionCarrera",
    "AvanceCurricular",
    "Facultad",
    "HorariosCatedra",
    "HorariosLaboratorio",
    "MotivoMinor",
    "ExpectativasMinor",
    "Comentarios"
]


SCHEDULE_TABLE_MARKER = "Horario Teoría nombre_asignatura"

def set_cell_border(cell, color="000000", size="12", space="0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), space)
        element.set(qn("w:color"), color)

def load_postulacion_dataframe(excel_path: str, sheet_name: str) -> pd.DataFrame:
    df = read_excel_file(excel_path, sheet_name)
    df = df.rename(columns=COLUMN_RENAME_MAP)
    validate_non_empty_dataframe(df)
    validate_required_columns(df, REQUIRED_COLUMNS)
    return df


def replace_in_paragraphs(doc, row_data: dict, config: dict):
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace("fecha_ingreso", config["fecha_documento"])
        paragraph.text = paragraph.text.replace("semestre_ingreso", config["semestre"])
        paragraph.text = paragraph.text.replace("nombre_jefe_carrera", safe_str(row_data["JefeCarrera"]))
        paragraph.text = paragraph.text.replace("carrera_estudiante", safe_str(row_data["Carrera"]))


from copy import deepcopy

def clone_row_after(table, row_idx: int):
    tr = table.rows[row_idx]._tr
    new_tr = deepcopy(tr)
    tr.addnext(new_tr)
    return table.rows[row_idx + 1]


def fill_schedule_table(doc, row_data: dict, config: dict):
    horarios_catedra = config["horarios_catedra"]
    horarios_lab = config["horarios_lab"]

    tabla = find_table_by_text(doc, SCHEDULE_TABLE_MARKER)
    if tabla is None:
        raise ValueError(
            "No se encontró la tabla de horarios en la plantilla. "
            "Verifica que exista el marcador 'Horario Teoría nombre_asignatura'."
        )

    cabecera_index = find_row_index_by_cell_text(tabla, SCHEDULE_TABLE_MARKER)
    if cabecera_index is None:
        raise ValueError("No se encontró la fila cabecera de horarios en la plantilla.")

    total_filas = max(len(horarios_catedra), len(horarios_lab))
    if total_filas == 0:
        return

    for i in range(total_filas):
        hora_catedra = safe_str(horarios_catedra[i]) if i < len(horarios_catedra) else ""
        hora_lab = safe_str(horarios_lab[i]) if i < len(horarios_lab) else ""

        nueva_fila = tabla.add_row()

        nueva_fila.cells[0].text = hora_catedra
        nueva_fila.cells[1].text = f"respuesta_catedra_{i+1}" if hora_catedra else ""
        nueva_fila.cells[2].text = hora_lab
        nueva_fila.cells[3].text = f"respuesta_lab_{i+1}" if hora_lab else ""

        for j in range(4):
            cell = nueva_fila.cells[j]

            if cell.paragraphs:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            set_cell_border(cell, color="000000", size="4")

def replace_in_tables(doc, row_data: dict, config: dict):
    horarios_catedra = config["horarios_catedra"]
    horarios_lab = config["horarios_lab"]
    nombre_asignatura = config["nombre_asignatura"]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "primer_nombresegundo_nombre" in cell.text:
                    cell.text = f"{safe_str(row_data['PrimerNombre'])} {safe_str(row_data['SegundoNombre'])}".strip()

                if "primer_apellidosegundo_apellido" in cell.text:
                    cell.text = f"{safe_str(row_data['ApellidoPaterno'])} {safe_str(row_data['ApellidoMaterno'])}".strip()

                if "rut_estudiante" in cell.text:
                    cell.text = safe_str(row_data["RUT"])

                if "correo_estudiante" in cell.text:
                    cell.text = safe_str(row_data["CorreoInstitucional"])

                if "carrera_estudiante" in cell.text:
                    cell.text = safe_str(row_data["Carrera"])

                if "facultad_estudiante" in cell.text:
                    cell.text = safe_str(row_data["Facultad"])

                if "duracion_carrera" in cell.text:
                    cell.text = safe_str(row_data["DuracionCarrera"])

                if "nivel_avance" in cell.text:
                    cell.text = safe_str(row_data["AvanceCurricular"])

                if "Horario Teoría nombre_asignatura" in cell.text:
                    cell.text = f"Horario Teoría {nombre_asignatura}"
                    if cell.paragraphs:
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].font.bold = False

                if "Horario Laboratorio nombre_asignatura" in cell.text:
                    cell.text = f"Horario Laboratorio {nombre_asignatura}"
                    if cell.paragraphs:
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for i in range(len(horarios_catedra)):
                    if f"respuesta_catedra_{i+1}" in cell.text:
                        if safe_str(row_data["HorariosCatedra"]) == safe_str(horarios_catedra[i]):
                            cell.text = "X"
                            if cell.paragraphs:
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            cell.text = ""
                        break

                for i in range(len(horarios_lab)):
                    if f"respuesta_lab_{i+1}" in cell.text:
                        if safe_str(row_data["HorariosLaboratorio"]) == safe_str(horarios_lab[i]):
                            cell.text = "X"
                            if cell.paragraphs:
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            cell.text = ""
                        break

                if "motivacion_minor" in cell.text:
                    cell.text = safe_str(row_data["MotivoMinor"])

                if "aprendizaje_minor" in cell.text:
                    cell.text = safe_str(row_data["ExpectativasMinor"])

                if "comentarios_estudiantes" in cell.text:
                    cell.text = safe_str(row_data["Comentarios"])


def insert_page_break_if_needed(doc):
    for i, paragraph in enumerate(doc.paragraphs):
        if "salto_pagina" in paragraph.text:
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i + 1].insert_paragraph_before("").add_run().add_break(WD_BREAK.PAGE)
            paragraph.clear()
            break


def generate_document(row_data: dict, save_path: str, config: dict):
    template_path = config["template_path"]
    validate_file_exists(template_path, "archivo de plantilla Word")

    doc = Document(template_path)

    fill_schedule_table(doc, row_data, config)
    replace_in_paragraphs(doc, row_data, config)
    replace_in_tables(doc, row_data, config)
    insert_page_break_if_needed(doc)

    try:
        doc._element.body.remove(doc._element.body[0])
    except Exception:
        pass

    doc.save(save_path)


def build_output_path(base_output_folder: str, row_data: dict) -> Path:
    carrera = sanitize_filename(safe_str(row_data["Carrera"]) or "SinCarrera")
    first_name = sanitize_filename(safe_str(row_data["PrimerNombre"]))
    last_name_1 = sanitize_filename(safe_str(row_data["ApellidoPaterno"]))
    last_name_2 = sanitize_filename(safe_str(row_data["ApellidoMaterno"]))

    folder = ensure_folder(Path(base_output_folder) / carrera)
    filename = f"formulario_{first_name}_{last_name_1}_{last_name_2}.docx"
    return folder / filename


def process_postulacion(excel_path: str, output_folder: str, config: dict, logger=None) -> dict:
    def log(msg: str):
        if logger:
            logger(msg)

    validate_file_exists(excel_path, "archivo Excel")
    validate_file_exists(config["template_path"], "archivo de plantilla Word")

    result = {
        "total": 0,
        "ok": 0,
        "errors": 0,
        "error_details": []
    }

    df = load_postulacion_dataframe(excel_path, config["sheet_name"])
    result["total"] = len(df)
    ensure_folder(output_folder)

    for idx, row in df.iterrows():
        row_data = row.to_dict()

        nombre = f"{safe_str(row_data.get('PrimerNombre', ''))} {safe_str(row_data.get('ApellidoPaterno', ''))}".strip()
        carrera = safe_str(row_data.get("Carrera", ""))

        try:
            log(f"Procesando fila {idx + 1}: {nombre} ({carrera})")
            save_path = build_output_path(output_folder, row_data)
            generate_document(row_data, str(save_path), config)
            result["ok"] += 1
            log(f"OK -> {save_path}")
        except Exception as e:
            result["errors"] += 1
            detail = f"Error en fila {idx + 1} ({nombre}): {e}"
            result["error_details"].append(detail)
            log(detail)

    return result