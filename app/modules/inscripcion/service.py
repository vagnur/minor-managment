from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.file_utils import ensure_folder, sanitize_filename, safe_str
from app.core.validation_utils import (
    validate_file_exists,
    validate_required_columns,
)
from app.core.docx_utils import find_table_by_text


SCHEDULE_TABLE_MARKER = "Horario Teoría nombre_asignatura"
TAAA_LAB_TABLE_MARKER = "Horario Laboratorio Taller de Aprendizaje Automático Aplicado"


def load_workbook_sheets(excel_path: str) -> dict[str, pd.DataFrame]:
    return pd.read_excel(excel_path, sheet_name=None)


def get_available_subjects(config: dict) -> list[str]:
    return list(config["subjects"].keys())


def normalize_subject_dataframe(df: pd.DataFrame, config: dict, subject_config: dict) -> pd.DataFrame:
    full_mapping = {}
    full_mapping.update(config["common_column_mapping"])
    full_mapping.update(subject_config["column_mapping"])

    df = df.rename(columns=full_mapping)

    catedra_field = subject_config.get("horarios_catedra_field", "")
    if catedra_field and catedra_field in df.columns:
        df["HorariosCatedra"] = df[catedra_field]

    lab_field = subject_config.get("horarios_lab_field", "")
    if lab_field and lab_field in df.columns:
        df["HorariosLaboratorio"] = df[lab_field]

    disponibles_field = subject_config.get("horarios_disponibles_field", "")
    if disponibles_field and disponibles_field in df.columns:
        df["HorariosDisponibles"] = df[disponibles_field]

    return df


def get_required_columns(subject_config: dict) -> list[str]:
    required = [
        "PrimerNombre",
        "ApellidoPaterno",
        "ApellidoMaterno",
        "RUT",
        "CorreoInstitucional",
        "Carrera",
        "JefeCarrera",
        "DuracionCarrera",
        "AvanceCurricular",
        "Facultad",
    ]

    if subject_config.get("has_catedra", False):
        required.append("HorariosCatedra")

    if subject_config.get("has_lab", False):
        required.append("HorariosLaboratorio")

    return required


def is_effectively_empty(df: pd.DataFrame) -> bool:
    if df.empty:
        return True

    temp = df.copy().dropna(how="all")
    return temp.empty


def validate_subject_dataframe(df: pd.DataFrame, subject_config: dict):
    validate_required_columns(df, get_required_columns(subject_config))


def build_output_path(base_output_folder: str, subject_name: str, row_data: dict) -> Path:
    subject_folder = sanitize_filename(subject_name)
    carrera = sanitize_filename(safe_str(row_data["Carrera"]) or "SinCarrera")
    first_name = sanitize_filename(safe_str(row_data["PrimerNombre"]))
    last_name_1 = sanitize_filename(safe_str(row_data["ApellidoPaterno"]))
    last_name_2 = sanitize_filename(safe_str(row_data["ApellidoMaterno"]))

    folder = ensure_folder(Path(base_output_folder) / subject_folder / carrera)
    filename = f"formulario_{first_name}_{last_name_1}_{last_name_2}.docx"
    return folder / filename


def set_cell_border(cell, color="000000", size="4", space="0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), space)
        element.set(qn("w:color"), color)


def replace_in_paragraphs(doc, row_data: dict, semestre: str, fecha_documento: str):
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace("fecha_ingreso", safe_str(fecha_documento))
        paragraph.text = paragraph.text.replace("semestre_ingreso", safe_str(semestre))
        paragraph.text = paragraph.text.replace("nombre_jefe_carrera", safe_str(row_data["JefeCarrera"]))
        paragraph.text = paragraph.text.replace("carrera_estudiante", safe_str(row_data["Carrera"]))


def fill_schedule_table_regular(doc, subject_config: dict):
    horarios_catedra = subject_config["horarios_catedra"]
    horarios_lab = subject_config["horarios_lab"]

    tabla = find_table_by_text(doc, SCHEDULE_TABLE_MARKER)
    if tabla is None:
        raise ValueError(
            "No se encontró la tabla de horarios en la plantilla general. "
            "Verifica que exista el marcador 'Horario Teoría nombre_asignatura'."
        )

    total_filas = max(len(horarios_catedra), len(horarios_lab))
    if total_filas == 0:
        return

    for i in range(total_filas):
        hora_catedra = safe_str(horarios_catedra[i]) if i < len(horarios_catedra) else ""
        hora_lab = safe_str(horarios_lab[i]) if i < len(horarios_lab) else ""

        nueva_fila = tabla.add_row()

        if len(nueva_fila.cells) < 4:
            raise ValueError("La fila agregada en la plantilla general no tiene al menos 4 celdas.")

        nueva_fila.cells[0].text = hora_catedra
        nueva_fila.cells[1].text = f"respuesta_catedra_{i+1}" if hora_catedra else ""
        nueva_fila.cells[2].text = hora_lab
        nueva_fila.cells[3].text = f"respuesta_lab_{i+1}" if hora_lab else ""

        for j in range(4):
            cell = nueva_fila.cells[j]
            if cell.paragraphs:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_cell_border(cell, color="000000", size="4")


def fill_schedule_table_taaa(doc, subject_config: dict):
    horarios_lab = subject_config["horarios_lab"]

    tabla = find_table_by_text(doc, TAAA_LAB_TABLE_MARKER)
    if tabla is None:
        raise ValueError(
            "No se encontró la tabla de horarios en la plantilla TAAA. "
            "Verifica que exista el marcador "
            "'Horario Laboratorio Taller de Aprendizaje Automático Aplicado'."
        )

    if len(horarios_lab) == 0:
        return

    for i, hora_lab in enumerate(horarios_lab):
        nueva_fila = tabla.add_row()

        if len(nueva_fila.cells) < 4:
            raise ValueError("La fila agregada en la plantilla TAAA no tiene al menos 4 celdas.")

        nueva_fila.cells[0].merge(nueva_fila.cells[1])
        nueva_fila.cells[2].merge(nueva_fila.cells[3])

        nueva_fila.cells[0].text = safe_str(hora_lab)
        nueva_fila.cells[2].text = f"respuesta_lab_{i+1}"

        for j in (0, 2):
            cell = nueva_fila.cells[j]
            if cell.paragraphs:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_cell_border(cell, color="000000", size="4")


def replace_in_tables_regular(doc, row_data: dict, subject_config: dict):
    horarios_catedra = subject_config["horarios_catedra"]
    horarios_lab = subject_config["horarios_lab"]
    display_name = subject_config.get("display_name", "")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text

                if "primer_nombresegundo_nombre" in text:
                    cell.text = f"{safe_str(row_data['PrimerNombre'])} {safe_str(row_data.get('SegundoNombre', ''))}".strip()

                if "primer_apellidosegundo_apellido" in text:
                    cell.text = f"{safe_str(row_data['ApellidoPaterno'])} {safe_str(row_data['ApellidoMaterno'])}".strip()

                if "rut_estudiante" in text:
                    cell.text = safe_str(row_data["RUT"])

                if "correo_estudiante" in text:
                    cell.text = safe_str(row_data["CorreoInstitucional"])

                if "carrera_estudiante" in text:
                    cell.text = safe_str(row_data["Carrera"])

                if "facultad_estudiante" in text:
                    cell.text = safe_str(row_data["Facultad"])

                if "duracion_carrera" in text:
                    cell.text = safe_str(row_data["DuracionCarrera"])

                if "nivel_avance" in text:
                    cell.text = safe_str(row_data["AvanceCurricular"])

                if "Horario Teoría nombre_asignatura" in text:
                    cell.text = f"Horario Teoría {display_name}"
                    if cell.paragraphs:
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                if "Horario Laboratorio nombre_asignatura" in text:
                    cell.text = f"Horario Laboratorio {display_name}"
                    if cell.paragraphs:
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for i in range(len(horarios_catedra)):
                    if f"respuesta_catedra_{i+1}" in cell.text:
                        if safe_str(row_data["HorariosCatedra"]) == safe_str(horarios_catedra[i]):
                            cell.text = "X"
                            if cell.paragraphs:
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            cell.text = ""
                        break

                for i in range(len(horarios_lab)):
                    if f"respuesta_lab_{i+1}" in cell.text:
                        if safe_str(row_data["HorariosLaboratorio"]) == safe_str(horarios_lab[i]):
                            cell.text = "X"
                            if cell.paragraphs:
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            cell.text = ""
                        break


def replace_in_tables_taaa(doc, row_data: dict, subject_config: dict):
    horarios_lab = subject_config["horarios_lab"]
    display_name = subject_config.get("display_name", TAAA_LAB_TABLE_MARKER)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text

                if "primer_nombresegundo_nombre" in text:
                    cell.text = f"{safe_str(row_data['PrimerNombre'])} {safe_str(row_data.get('SegundoNombre', ''))}".strip()

                if "primer_apellidosegundo_apellido" in text:
                    cell.text = f"{safe_str(row_data['ApellidoPaterno'])} {safe_str(row_data['ApellidoMaterno'])}".strip()

                if "rut_estudiante" in text:
                    cell.text = safe_str(row_data["RUT"])

                if "correo_estudiante" in text:
                    cell.text = safe_str(row_data["CorreoInstitucional"])

                if "carrera_estudiante" in text:
                    cell.text = safe_str(row_data["Carrera"])

                if "facultad_estudiante" in text:
                    cell.text = safe_str(row_data["Facultad"])

                if "duracion_carrera" in text:
                    cell.text = safe_str(row_data["DuracionCarrera"])

                if "nivel_avance" in text:
                    cell.text = safe_str(row_data["AvanceCurricular"])

                if TAAA_LAB_TABLE_MARKER in text:
                    cell.text = f"Horario Laboratorio {display_name}"
                    if cell.paragraphs:
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for i in range(len(horarios_lab)):
                    if f"respuesta_lab_{i+1}" in cell.text:
                        if safe_str(row_data["HorariosLaboratorio"]) == safe_str(horarios_lab[i]):
                            cell.text = "X"
                            if cell.paragraphs:
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            cell.text = ""
                        break


def insert_page_break_if_needed(doc):
    for i, paragraph in enumerate(doc.paragraphs):
        if "salto_pagina" in paragraph.text:
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i + 1].insert_paragraph_before("").add_run().add_break(WD_BREAK.PAGE)
            paragraph.clear()
            break


def generate_document_for_row(
    row_data: dict,
    subject_config: dict,
    semestre: str,
    fecha_documento: str,
    save_path: str,
):
    template_path = subject_config["template_path"]
    doc = Document(template_path)

    if subject_config.get("has_catedra", False):
        fill_schedule_table_regular(doc, subject_config)
        replace_in_tables_regular(doc, row_data, subject_config)
    else:
        fill_schedule_table_taaa(doc, subject_config)
        replace_in_tables_taaa(doc, row_data, subject_config)

    replace_in_paragraphs(doc, row_data, semestre, fecha_documento)
    insert_page_break_if_needed(doc)

    try:
        doc._element.body.remove(doc._element.body[0])
    except Exception:
        pass

    doc.save(save_path)


def validate_subject_resources(subject_name: str, subject_config: dict, runtime_config: dict):
    validate_file_exists(subject_config["template_path"], f"plantilla de {subject_name}")

    if subject_config.get("has_catedra", False) and not runtime_config.get("horarios_catedra", []):
        raise ValueError(f"La asignatura {subject_name} no tiene horarios de cátedra ingresados.")

    if subject_config.get("has_lab", False) and not runtime_config.get("horarios_lab", []):
        raise ValueError(f"La asignatura {subject_name} no tiene horarios de laboratorio ingresados.")


def process_subject(
    subject_name: str,
    df: pd.DataFrame,
    config: dict,
    runtime_config: dict,
    output_folder: str,
    semestre: str,
    fecha_documento: str,
    logger=None,
) -> dict:
    def log(msg: str):
        if logger:
            logger(msg)

    result = {
        "subject": subject_name,
        "total": 0,
        "ok": 0,
        "errors": 0,
        "skipped": False,
        "error_details": [],
    }

    subject_config = config["subjects"][subject_name]
    subject_config = {
        **subject_config,
        "horarios_catedra": runtime_config.get("horarios_catedra", []),
        "horarios_lab": runtime_config.get("horarios_lab", []),
    }

    validate_subject_resources(subject_name, subject_config, runtime_config)

    if is_effectively_empty(df):
        result["skipped"] = True
        log(f"[{subject_name}] Hoja vacía. Se omite.")
        return result

    df = normalize_subject_dataframe(df, config, subject_config)
    validate_subject_dataframe(df, subject_config)
    df = df.dropna(how="all")

    result["total"] = len(df)

    log(f"[{subject_name}] Registros a procesar: {len(df)}")

    for idx, row in df.iterrows():
        row_data = row.to_dict()

        try:
            save_path = build_output_path(output_folder, subject_name, row_data)
            generate_document_for_row(
                row_data=row_data,
                subject_config=subject_config,
                semestre=semestre,
                fecha_documento=fecha_documento,
                save_path=str(save_path),
            )
            result["ok"] += 1
            log(f"[{subject_name}] OK fila {idx + 1}: {save_path}")
        except Exception as e:
            result["errors"] += 1
            detail = f"[{subject_name}] Error fila {idx + 1}: {e}"
            result["error_details"].append(detail)
            log(detail)

    return result


def validate_excel_workbook(excel_path: str, config: dict, selected_subjects: list[str]) -> dict:
    validate_file_exists(excel_path, "archivo Excel")

    workbook = load_workbook_sheets(excel_path)
    results = {
        "available_sheets": list(workbook.keys()),
        "missing_sheets": [],
        "empty_sheets": [],
    }

    for subject_name in selected_subjects:
        sheet_name = config["subjects"][subject_name]["sheet_name"]

        if sheet_name not in workbook:
            results["missing_sheets"].append(sheet_name)
            continue

        df = workbook[sheet_name]
        if is_effectively_empty(df):
            results["empty_sheets"].append(sheet_name)

    return results


def process_inscripcion(
    excel_path: str,
    output_folder: str,
    semestre: str,
    fecha_documento: str,
    selected_subjects: list[str],
    subject_runtime_configs: dict,
    config: dict,
    logger=None,
) -> dict:
    def log(msg: str):
        if logger:
            logger(msg)

    validate_file_exists(excel_path, "archivo Excel")
    ensure_folder(output_folder)

    workbook = load_workbook_sheets(excel_path)

    global_result = {
        "subjects_processed": [],
        "subjects_skipped": [],
        "missing_sheets": [],
        "total_ok": 0,
        "total_errors": 0,
        "details": [],
    }

    for subject_name in selected_subjects:
        sheet_name = config["subjects"][subject_name]["sheet_name"]

        if sheet_name not in workbook:
            global_result["missing_sheets"].append(sheet_name)
            log(f"[{subject_name}] Hoja no encontrada en Excel: {sheet_name}")
            continue

        df = workbook[sheet_name]

        try:
            result = process_subject(
                subject_name=subject_name,
                df=df,
                config=config,
                runtime_config=subject_runtime_configs.get(subject_name, {}),
                output_folder=output_folder,
                semestre=semestre,
                fecha_documento=fecha_documento,
                logger=log,
            )

            if result["skipped"]:
                global_result["subjects_skipped"].append(subject_name)
            else:
                global_result["subjects_processed"].append(subject_name)

            global_result["total_ok"] += result["ok"]
            global_result["total_errors"] += result["errors"]
            global_result["details"].append(result)

        except Exception as e:
            detail = f"[{subject_name}] Error general de asignatura: {e}"
            global_result["total_errors"] += 1
            global_result["details"].append({
                "subject": subject_name,
                "total": 0,
                "ok": 0,
                "errors": 1,
                "skipped": False,
                "error_details": [detail],
            })
            log(detail)

    return global_result