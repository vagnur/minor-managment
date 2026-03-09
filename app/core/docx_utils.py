from copy import deepcopy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def find_table_by_headers(doc, expected_headers: list[str]):
    normalized_expected = [h.strip().upper() for h in expected_headers]

    for table in doc.tables:
        if not table.rows:
            continue

        first_row = table.rows[0]
        row_headers = [cell.text.strip().upper() for cell in first_row.cells]

        if row_headers[:len(normalized_expected)] == normalized_expected:
            return table

    return None


def replace_text_in_paragraphs(doc, replacements: dict[str, str]):
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)


def replace_text_in_tables(doc, replacements: dict[str, str]):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in replacements.items():
                    if old in cell.text:
                        cell.text = cell.text.replace(old, new)


def clone_last_row(table):
    """
    Clona la última fila de la tabla conservando mejor el formato.
    """
    last_row = table.rows[-1]
    new_tr = deepcopy(last_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def clear_row_text(row):
    for cell in row.cells:
        cell.text = ""


def set_cell_no_wrap(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap)


def set_table_fixed_layout(table):
    """
    Reduce quiebres raros al impedir autoajuste agresivo.
    """
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

def find_table_by_text(doc, target_text: str):
    """
    Busca una tabla que contenga un texto específico en cualquiera de sus celdas.
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if target_text in cell.text:
                    return table
    return None


def find_row_index_by_cell_text(table, target_text: str):
    """
    Busca el índice de la fila que contiene el texto indicado.
    """
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if target_text in cell.text:
                return i
    return None

def find_table_by_headers(doc, expected_headers: list[str]):
    """
    Busca una tabla cuyos encabezados contengan exactamente los textos esperados.
    """
    normalized_expected = [h.strip().upper() for h in expected_headers]

    for table in doc.tables:
        if not table.rows:
            continue

        first_row = table.rows[0]
        row_headers = [cell.text.strip().upper() for cell in first_row.cells]

        if row_headers[:len(normalized_expected)] == normalized_expected:
            return table

    return None


def replace_text_in_paragraphs(doc, replacements: dict[str, str]):
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)


def replace_text_in_tables(doc, replacements: dict[str, str]):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in replacements.items():
                    if old in cell.text:
                        cell.text = cell.text.replace(old, new)